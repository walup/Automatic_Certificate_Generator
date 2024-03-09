import qrcode
from qrcode.image.pure import PyPNGImage
import pandas as pd
import numpy as np 
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Cm
from docx2pdf import convert


class Text2QR:

    def __init__(self, labels, data, boxSize = 10, border = 2):

        self.labels = labels
        self.data = data

        self.boxSize = boxSize
        self.border = border

    
    def saveQR(self, pathName):
        '''
        Saves a QR image in the specified folder
        '''

        strQR = ""
        for i in range(0,len(self.labels)):
            if(len(self.data[i]) == 1):
                strQR = strQR + self.labels[i] + ": "+str(self.data[i][0])+ "\n"
            elif(len(self.data[i]) > 1):
                strQR = strQR + self.labels[i] + ": "+"\n"+"-----------------------"+"\n"
                for j in range(0,len(self.data[i])):
                    strQR = strQR + str(self.data[i][j])+"\n"
                strQR = strQR + "-----------------------"+"\n"
                
        
        qrCode = qrcode.QRCode(version = 1, error_correction = qrcode.constants.ERROR_CORRECT_L, box_size = self.boxSize, border = self.border)
        qrCode.add_data(strQR)
        qrCode.make(fit = True)
        img = qrCode.make_image()
        img.save(pathName)
    

    def renderQR(self):
        '''
        Generates a QR and returns its corresponding image object
        '''

        strQR = ""
        for i in range(0,len(self.labels)):
            if(len(self.data[i]) == 1):
                strQR = strQR + self.labels[i] + ": "+str(self.data[i][0])+ "\n"
            elif(len(self.data[i]) > 1):
                strQR = strQR + self.labels[i] + ": "+"\n"+"-----------------------"+"\n"
                for j in range(0,len(self.data[i])):
                    strQR = strQR + "-"+str(self.data[i][j])+"\n"
                
                strQR = strQR + "-----------------------"+"\n"
                
        
        qrCode = qrcode.QRCode(version = 1, error_correction = qrcode.constants.ERROR_CORRECT_L, box_size = self.boxSize, border = self.border)
        qrCode.add_data(strQR)
        qrCode.make(fit = True)
        img = qrCode.make_image()
        return img
    

class ExcelHandler:

    def __init__(self, path):
        self.path = path

    def getNEntries(self,sheetName = "Hoja1"):
        return len(pd.read_excel(self.path, sheet_name = sheetName).index)

    def getSheetNames(self):
        '''
        Returns a list of the excel sheet names
        '''

        file = pd.ExcelFile(self.path)
        names = file.sheet_names

        return names
    
    def getSheetField(self, field, sheetName = "Hoja1"):
        '''
        Returns the column data of a specified sheet and field (header)
        '''

        file = pd.read_excel(self.path, sheet_name = sheetName)
        dataFrame = file.loc[:,field]

        return list(dataFrame)
    
    def getSheetFieldNames(self, sheetName = "Hoja1"):
        file = pd.read_excel(self.path, sheet_name = sheetName)
        fieldNames = file.head()

        return list(fieldNames)
    
class WordEditor:

    def __init__(self, path):
        self.path = path
        self.document = Document(self.path)
    
    def substitutePattern(self, pattern, substitution):
        for p in self.document.paragraphs:
            if pattern in p.text:
                inLine = p.runs

                for i in range(len(inLine)):
                    if pattern in inLine[i].text:
                        text = inLine[i].text.replace(pattern, substitution)
                        inLine[i].text = text
    

    def insertImage(self, imageURL, width, height):
        paragraph = self.document.add_paragraph()
        #Center alignment
        paragraph.alignment = 1
        r = paragraph.add_run()
        r.add_picture(imageURL, width = Cm(width), height = Cm(height))

    def save(self, fileName):
        self.document.save(fileName)
    
    def saveAsPDF(self, fileNameWord, fileNamePDF):
        self.save(fileNameWord)
        inputFile = fileNameWord
        outputFile = fileNamePDF
        file = open(outputFile, "w")
        file.close()
        convert(inputFile, outputFile)








        

            
